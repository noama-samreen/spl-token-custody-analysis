import os
import json
from datetime import datetime
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

def create_styles():
    styles = getSampleStyleSheet()
    
    # Title style with better spacing and alignment
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=20,
        spaceAfter=30,
        alignment=1,  # Center alignment
        textColor=colors.black,
        fontName='Helvetica-Bold'
    )
    
    # Cell style with better formatting
    cell_style = ParagraphStyle(
        'CustomCell',
        parent=styles['Normal'],
        fontSize=10,
        spaceAfter=10,
        fontName='Helvetica',
        leading=14  # Line spacing
    )
    
    # Context style with better readability
    context_style = ParagraphStyle(
        'CustomContext',
        parent=styles['Normal'],
        fontSize=11,
        spaceAfter=12,
        leading=16,  # Increased line spacing
        alignment=4,  # Justified text
        fontName='Helvetica'
    )
    
    # Header style for table headers
    header_style = ParagraphStyle(
        'CustomHeader',
        parent=styles['Normal'],
        fontSize=11,
        fontName='Helvetica-Bold',
        textColor=colors.black,
        alignment=0  # Left alignment
    )
    
    return styles, title_style, cell_style, context_style, header_style

def create_basic_table(data, cell_style):
    """Create and style the basic information table"""
    # Reduced table width (adjusted from 6 inches to 5 inches total)
    table = Table(data, colWidths=[1.2*inch, 3.8*inch])
    table.setStyle(TableStyle([
        ('ALIGN', (0,0), (-1,-1), 'LEFT'),
        ('FONTNAME', (0,0), (-1,-1), 'Helvetica'),
        ('FONTSIZE', (0,0), (-1,-1), 10),
        ('GRID', (0,0), (-1,-1), 1, colors.black),
        ('BACKGROUND', (0,0), (0,-1), colors.lightgrey),
        ('TEXTCOLOR', (0,0), (-1,-1), colors.black),
        ('PADDING', (0,0), (-1,-1), 6),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
    ]))
    return table

def create_additional_table(data, cell_style):
    """Create and style the additional fields table"""
    table = Table(data, colWidths=[2.5*inch, 3.5*inch])
    table.setStyle(TableStyle([
        ('ALIGN', (0,0), (-1,-1), 'LEFT'),
        ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),  # Bold header row
        ('FONTNAME', (0,1), (-1,-1), 'Helvetica'),
        ('FONTSIZE', (0,0), (-1,-1), 10),
        ('GRID', (0,0), (-1,-1), 1, colors.black),
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#f0f0f0')),  # Light gray header
        ('TEXTCOLOR', (0,0), (-1,-1), colors.black),
        ('PADDING', (0,0), (-1,-1), 12),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ('ROWBACKGROUNDS', (0,1), (-1,-2), [colors.white, colors.HexColor('#f9f9f9')]),  # Alternating rows except last
        ('BACKGROUND', (0,-1), (-1,-1), colors.HexColor('#f8f8f8')),  # Slight emphasis on last row
    ]))
    return table

def create_pdf(token_data, output_dir):
    """Generate PDF for a single token"""
    # Handle missing or invalid name/symbol
    token_name = token_data.get('name', 'Unknown')
    if token_name in ['N/A', None, '']:
        token_name = 'Unknown'
    
    token_symbol = token_data.get('symbol', 'UNKNOWN')
    if token_symbol in ['N/A', None, '']:
        token_symbol = 'UNKNOWN'
    
    filename = f"{token_name} ({token_symbol}) Security Memo.pdf"
    # Sanitize filename to remove invalid characters
    filename = "".join(c for c in filename if c.isalnum() or c in (' ', '-', '_', '(', ')', '.'))
    filepath = os.path.join(output_dir, filename)
    
    doc = SimpleDocTemplate(
        filepath,
        pagesize=letter,
        leftMargin=72,
        rightMargin=72,
        topMargin=72,
        bottomMargin=72
    )
    
    styles, title_style, cell_style, context_style, header_style = create_styles()
    elements = []
    
    # Title with better spacing and error handling
    title = Paragraph(
        f"Solana Token Security Assessment:<br/>{token_name} ({token_symbol})", 
        title_style
    )
    elements.append(title)
    elements.append(Spacer(1, 30))
    
    # Basic information table with wrapped text
    current_date = datetime.now().strftime("%Y-%m-%d")
    profile = "SPL Token 2022 Standard" if "Token 2022" in token_data['owner_program'] else "SPL Token Standard"
    
    data = [
        [Paragraph("Reviewer", cell_style), Paragraph("Noama Samreen", cell_style)],
        [Paragraph("Profile", cell_style), Paragraph(profile, cell_style)],
        [Paragraph("Review Date", cell_style), Paragraph(current_date, cell_style)],
        [Paragraph("Network", cell_style), Paragraph("Solana", cell_style)],
        [Paragraph("Address", cell_style), Paragraph(token_data['address'], cell_style)]
    ]
    
    elements.append(create_basic_table(data, cell_style))
    elements.append(Spacer(1, 30))
    
    # Context text
    context_text = """<b>Solana SPL Token Review Context:</b> Solana tokens do not possess customizable code per 
asset. Rather, a single "program" generates boiler template tokens with distinct states for each 
newly created token. Therefore, examining the base program configurations is adequate for 
reviewing all other tokens associated with it. The 'Token Program' adheres to standard 
practices, undergoing thorough review and auditing procedures. Therefore, within this review 
process, the focus remains on validating token configurations specific to tokens managed by the 
trusted Token Program"""
    
    elements.append(Paragraph(context_text, context_style))
    elements.append(Spacer(1, 25))
    
    # Recommendation with error handling
    security_review = token_data.get('security_review', 'UNKNOWN')
    if security_review in ['N/A', None, '']:
        security_review = 'UNKNOWN'
    
    recommendation = (
        f"<b>{token_name} ({token_symbol}) "
        f"{'is' if security_review == 'PASSED' else 'is not'} recommended for listing.</b>"
    )
    elements.append(Paragraph(recommendation, ParagraphStyle(
        'CustomRecommendation',
        parent=context_style,
        fontSize=12,
        textColor=colors.HexColor('#006400') if security_review == 'PASSED' else colors.red
    )))
    elements.append(Spacer(1, 25))
    
    # Additional details table with error handling
    additional_data = [["Field", "Value"]]
    
    # Base fields for all tokens
    field_order = [
        'owner_program',
        'freeze_authority',
        'update_authority'
    ]
    
    # Add Token 2022 specific fields only if it's a Token 2022 program
    if "Token 2022" in token_data.get('owner_program', ''):
        field_order.extend([
            'permanent_delegate',
            'transaction_fees',
            'transfer_hook',
            'confidential_transfers'
        ])
    
    # Add pump.fun specific fields if update authority matches
    if "Pump.Fun Mint Authority" in str(token_data.get('update_authority', '')):
        field_order.extend([
            'is_genuine_pump_fun_token',
            'interacted_with',
            'token_graduated_to_raydium'
        ])
        # Only add these fields if there was an interaction
        if token_data.get('interacting_account') or token_data.get('interaction_signature'):
            field_order.extend([
                'interacting_account',
                'interaction_signature'
            ])
    
    # Add fields in specified order with error handling
    for field in field_order:
        value = token_data.get(field, 'None')
        if value in ['N/A', None, '']:
            value = 'None'
            
        # Special handling for owner program to show address and name
        if field == 'owner_program' and value != 'None':
            # The value already includes the program name from the JSON
            pass
        
        # Format boolean values
        if isinstance(value, bool):
            value = str(value)
            
        # Format the field name for display
        display_name = str(field).replace('_', ' ').title()
        
        additional_data.append([
            Paragraph(display_name, cell_style),
            Paragraph(str(value), cell_style)
        ])
    
    # Add security review as the last row
    security_style = ParagraphStyle(
        'SecurityCell',
        parent=cell_style,
        textColor=colors.HexColor('#006400') if security_review == 'PASSED' 
                 else colors.red if security_review == 'FAILED'
                 else colors.black,
        fontName='Helvetica-Bold'
    )
    
    additional_data.append([
        Paragraph("Security Review", cell_style),
        Paragraph(security_review, security_style)
    ])
    
    elements.append(create_additional_table(additional_data, cell_style))
    
    # Build PDF
    doc.build(elements)
    return filepath

def create_doc(token_data, output_dir):
    """Generate DOCX for a single token"""
    # Handle missing or invalid name/symbol
    token_name = token_data.get('name', 'Unknown')
    if token_name in ['N/A', None, '']:
        token_name = 'Unknown'
    
    token_symbol = token_data.get('symbol', 'UNKNOWN')
    if token_symbol in ['N/A', None, '']:
        token_symbol = 'UNKNOWN'
    
    filename = f"{token_name} ({token_symbol}) Security Memo.docx"
    filename = "".join(c for c in filename if c.isalnum() or c in (' ', '-', '_', '(', ')', '.'))
    filepath = os.path.join(output_dir, filename)
    
    doc = Document()
    
    # Add confidentiality notice
    confidentiality = doc.add_paragraph()
    confidentiality.add_run("Confidential treatment requested under NY Banking Law § 36.10 and NY Pub. Off. Law § 87.2(d).").italic = True
    doc.add_paragraph()
    
    # Title
    title = doc.add_heading(f"Solana Token Security Assessment:\n{token_name} ({token_symbol})", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    # Conflicts certification
    conflicts = doc.add_paragraph()
    conflicts.add_run("Conflicts Certification: ").bold = True
    conflicts.add_run("To the best of your knowledge, please confirm that you and your immediate family: (1) have not invested more than $10,000 in the asset or its issuer, (2) do not own more than 1% of the asset outstanding, and (3) do not have a personal relationship with the issuer's management, governing body, or owners. For wrapped assets, the underlying asset must be considered for the purpose of this conflict certification, unless: 1) the asset is a stablecoin; or 2) has a market cap of over $100 billion dollars. For multi-chain assets every version of the multi-chain asset must be counted together for the purpose of this conflict certification.")
    doc.add_paragraph()
    
    # Basic information table
    current_date = datetime.now().strftime("%Y-%m-%d")
    profile = "SPL Token 2022 Standard" if "Token 2022" in token_data['owner_program'] else "SPL Token Standard"
    
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    
    data = [
        ["Reviewer", "Noama Samreen"],
        ["Profile", profile],
        ["Review Date", current_date],
        ["Network", "Solana"],
        ["Address", token_data['address']]
    ]
    
    for i, (key, value) in enumerate(data):
        row = table.rows[i]
        row.cells[0].text = key
        row.cells[1].text = value
    
    doc.add_paragraph()
    
    # Context
    context = doc.add_paragraph()
    context.add_run("Solana SPL Token Review Context: ").bold = True
    context.add_run("Solana tokens do not possess customizable code per asset. Rather, a single "program" generates boiler template tokens with distinct states for each newly created token. Therefore, examining the base program configurations is adequate for reviewing all other tokens associated with it. The 'Token Program' adheres to standard practices, undergoing thorough review and auditing procedures. Therefore, within this review process, the focus remains on validating token configurations specific to tokens managed by the trusted Token Program")
    doc.add_paragraph()
    
    # Recommendation
    security_review = token_data.get('security_review', 'UNKNOWN')
    if security_review in ['N/A', None, '']:
        security_review = 'UNKNOWN'
    
    recommendation = doc.add_paragraph()
    rec_text = f"{token_name} ({token_symbol}) {'is' if security_review == 'PASSED' else 'is not'} recommended for listing."
    recommendation.add_run(rec_text).bold = True
    doc.add_paragraph()
    
    # Risk scores
    risk_table = doc.add_table(rows=2, cols=2)
    risk_table.style = 'Table Grid'
    
    risk_data = [
        ["Residual Security Risk Score:", str(token_data.get('residual_risk_score', 'N/A'))],
        ["Inherent Security Risk Score:", str(token_data.get('inherent_risk_score', 'N/A'))]
    ]
    
    for i, (key, value) in enumerate(risk_data):
        row = risk_table.rows[i]
        row.cells[0].text = key
        row.cells[1].text = value
    
    doc.add_paragraph()
    
    # Additional details table
    details_table = doc.add_table(rows=1, cols=2)
    details_table.style = 'Table Grid'
    header_cells = details_table.rows[0].cells
    header_cells[0].text = 'Field'
    header_cells[1].text = 'Value'
    
    field_order = [
        'owner_program',
        'freeze_authority',
        'update_authority'
    ]
    
    if "Token 2022" in token_data.get('owner_program', ''):
        field_order.extend([
            'permanent_delegate',
            'transaction_fees',
            'transfer_hook',
            'confidential_transfers'
        ])
    
    if "Pump.Fun Mint Authority" in str(token_data.get('update_authority', '')):
        field_order.extend([
            'is_genuine_pump_fun_token',
            'interacted_with',
            'token_graduated_to_raydium'
        ])
        if token_data.get('interacting_account') or token_data.get('interaction_signature'):
            field_order.extend([
                'interacting_account',
                'interaction_signature'
            ])
    
    for field in field_order:
        value = token_data.get(field, 'None')
        if value in ['N/A', None, '']:
            value = 'None'
        if isinstance(value, bool):
            value = str(value)
        display_name = str(field).replace('_', ' ').title()
        row_cells = details_table.add_row().cells
        row_cells[0].text = display_name
        row_cells[1].text = str(value)
    
    # Add security review as last row
    row_cells = details_table.add_row().cells
    row_cells[0].text = 'Security Review'
    row_cells[1].text = security_review
    
    doc.save(filepath)
    return filepath

def generate_reports(token_data, output_dir):
    """Generate both PDF and DOCX reports"""
    pdf_path = create_pdf(token_data, output_dir)
    doc_path = create_doc(token_data, output_dir)
    return pdf_path, doc_path

if __name__ == '__main__':
    pdf_path, doc_path = generate_reports(token_data, output_dir)
